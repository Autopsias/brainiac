"""DOCX handler — python-docx. Paragraphs + tables (Markdown tables, headers
retained per HARDENED:grill — see .tables.rows_to_markdown)."""
from __future__ import annotations

from pathlib import Path
from typing import Any

from .base import ExtractResult, Handler, density_gate
from .tables import rows_to_markdown

try:
    import docx  # python-docx
    _HAS_DOCX = True
except ImportError:  # pragma: no cover
    _HAS_DOCX = False

MAX_DOCX_BYTES = 100 * 1024 * 1024


def _open_document(path: Path) -> Any | ExtractResult:
    try:
        size = path.stat().st_size
    except OSError:
        size = 0
    if size > MAX_DOCX_BYTES:
        return ExtractResult.quarantine("file_too_large")
    try:
        return docx.Document(str(path))
    except Exception as exc:
        return ExtractResult.quarantine(
            "docx_extraction_error",
            warnings=[f"{type(exc).__name__}: {exc}"],
        )


def _render_document(document: Any) -> ExtractResult:
    try:
        parts = _document_parts(document)
    except Exception as exc:
        return ExtractResult.quarantine(
            "docx_extraction_error",
            warnings=[f"{type(exc).__name__}: {exc}"],
        )
    body_md = "\n".join(part for part in parts if part.strip())
    reason = density_gate(body_md)
    if reason:
        return ExtractResult.quarantine(reason)
    return ExtractResult(markdown=body_md, metadata={"tables": len(document.tables)})


def _document_parts(document: Any) -> list[str]:
    """Render paragraphs and tables in their original body order."""
    parts: list[str] = []
    body = document.element.body
    table_by_element = {table._tbl: table for table in document.tables}
    paragraph_by_element = {paragraph._p: paragraph for paragraph in document.paragraphs}
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            rendered = _render_paragraph(paragraph_by_element.get(child))
        elif tag == "tbl":
            rendered = _render_table(table_by_element.get(child))
        else:
            rendered = ""
        if rendered:
            parts.append(rendered)
    return parts


def _render_paragraph(paragraph: Any) -> str:
    if paragraph is None:
        return ""
    text = paragraph.text.strip()
    if not text:
        return ""
    if paragraph.style and paragraph.style.name and paragraph.style.name.startswith("Heading"):
        return f"## {text}\n"
    return f"{text}\n"


def _render_table(table: Any) -> str:
    if table is None:
        return ""
    rows = [[cell.text for cell in row.cells] for row in table.rows]
    return rows_to_markdown(rows)


class DocxHandler(Handler):
    extensions = (".docx",)
    dependency_name = "python-docx"

    @classmethod
    def available(cls) -> bool:
        return _HAS_DOCX

    @classmethod
    def extract(cls, path: Path) -> ExtractResult:
        if not _HAS_DOCX:
            return ExtractResult.quarantine("missing_dependency:python-docx")
        document = _open_document(path)
        if isinstance(document, ExtractResult):
            return document
        return _render_document(document)
